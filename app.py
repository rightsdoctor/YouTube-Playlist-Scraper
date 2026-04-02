import streamlit as st
import json
import re
import os
import glob
import subprocess
import shutil
import zipfile
import threading
import time
import base64
import pandas as pd
from datetime import datetime
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

# ============================================================
# 페이지 설정
# ============================================================
st.set_page_config(page_title="YT Playlist Scraper", layout="wide")
st.title("YouTube Playlist Scraper")
st.caption("플레이리스트 URL → 메타데이터 + 자막 → Excel / CSV / 자막 파일")

# ============================================================
# 상수 & 경로
# ============================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SUBTITLE_DIR = os.path.join(BASE_DIR, "subtitles_temp")
CONVERTED_DIR = os.path.join(BASE_DIR, "subtitles_converted")
SUB_EXTS = ["srt", "vtt", "srv1", "srv2", "srv3", "ttml", "ass", "json3", "lrc"]


# ============================================================
# 헬퍼 함수
# ============================================================
def run_cmd(args, timeout=120):
    try:
        r = subprocess.run(args, capture_output=True, text=True, timeout=timeout)
        return r.stdout, r.stderr, r.returncode
    except subprocess.TimeoutExpired:
        return "", "TIMEOUT", -1
    except Exception as e:
        return "", str(e), -1


def check_ffmpeg():
    out, err, code = run_cmd(["ffmpeg", "-version"], timeout=5)
    return code == 0


def check_ytdlp():
    out, err, code = run_cmd(["yt-dlp", "--version"], timeout=10)
    return out.strip(), code == 0


def srt_to_plain_text(content: str) -> str:
    lines = content.strip().split('\n')
    text_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\d+$', line):
            continue
        if re.match(r'\d{2}:\d{2}:\d{2}', line):
            continue
        if line.startswith("WEBVTT") or line.startswith("Kind:") or line.startswith("Language:"):
            continue
        line = re.sub(r'<[^>]+>', '', line)
        line = re.sub(r'\{[^}]+\}', '', line)
        if line:
            text_lines.append(line)
    deduplicated = []
    for t in text_lines:
        if not deduplicated or t != deduplicated[-1]:
            deduplicated.append(t)
    return ' '.join(deduplicated)


def format_duration(seconds):
    if not seconds:
        return ''
    seconds = int(seconds)
    h, remainder = divmod(seconds, 3600)
    m, s = divmod(remainder, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


def find_subtitle_files_for_video(video_id, subtitle_dir):
    results = []
    if not os.path.exists(subtitle_dir):
        return results
    for fname in os.listdir(subtitle_dir):
        if fname.startswith(video_id + "."):
            fpath = os.path.join(subtitle_dir, fname)
            if os.path.isfile(fpath) and os.path.getsize(fpath) > 0:
                results.append(fpath)
    return results


def parse_subtitle_lang(filepath):
    fname = os.path.basename(filepath)
    parts = fname.rsplit('.', 2)
    if len(parts) == 3:
        return parts[1]
    return "unknown"


def read_subtitles_for_video(video_id, subtitle_dir):
    result = {}
    for fpath in find_subtitle_files_for_video(video_id, subtitle_dir):
        lang = parse_subtitle_lang(fpath)
        try:
            with open(fpath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read().strip()
            if content:
                result[lang] = content
        except Exception:
            pass
    return result


def count_videos_with_subs(subtitle_dir):
    vids = set()
    if not os.path.exists(subtitle_dir):
        return vids, 0
    total_files = 0
    for fname in os.listdir(subtitle_dir):
        fpath = os.path.join(subtitle_dir, fname)
        if os.path.isfile(fpath) and os.path.getsize(fpath) > 0:
            vid = fname.split('.')[0]
            if vid:
                vids.add(vid)
                total_files += 1
    return vids, total_files


def zip_directory_all(dir_path, ext):
    buf = BytesIO()
    matched = glob.glob(os.path.join(dir_path, f"*.{ext}"))
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for fpath in matched:
            zf.write(fpath, os.path.basename(fpath))
    return buf.getvalue(), len(matched)


def make_download_link(data: bytes, filename: str, label: str) -> str:
    b64 = base64.b64encode(data).decode()
    return (
        f'<a href="data:application/octet-stream;base64,{b64}" '
        f'download="{filename}" '
        f'style="display:inline-block;padding:0.5rem 1rem;'
        f'background-color:#FF4B4B;color:white;text-decoration:none;'
        f'border-radius:0.5rem;font-weight:600;text-align:center;'
        f'width:100%;box-sizing:border-box;">'
        f'{label}</a>'
    )


# ============================================================
# 사이드바
# ============================================================
with st.sidebar:
    st.header("설정")
    playlist_url = st.text_input(
        "플레이리스트 URL",
        placeholder="https://www.youtube.com/playlist?list=..."
    )

    st.subheader("자막 옵션")
    sub_mode = st.radio(
        "자막 수집 방식",
        ["수동 자막만", "자동 생성 자막만", "수동 우선, 없으면 자동", "수동 + 자동 모두"],
        index=2,
    )
    sub_mode_map = {
        "수동 자막만": "manual_only",
        "자동 생성 자막만": "auto_only",
        "수동 우선, 없으면 자동": "manual_first",
        "수동 + 자동 모두": "both",
    }
    sub_choice = sub_mode_map[sub_mode]
    sub_lang = st.text_input("자막 언어", value="ko",
                             help="예: ko, en, ja 또는 all")
    output_format = st.selectbox("자막 파일 포맷", ["txt", "srt", "vtt", "docx"])

    st.subheader("속도 조절")
    workers = st.slider(
        "동시 처리 수",
        min_value=1, max_value=10, value=1,
        help="1 = 순차 처리 (가장 안전), 높을수록 빠르지만 차단 위험 증가"
    )
    sleep_sec = st.slider(
        "요청 간 대기 (초)",
        min_value=0, max_value=15, value=3,
        help="각 영상 처리 후 대기 시간. 3~5초 권장"
    )
    max_consecutive_fails = st.slider(
        "연속 실패 시 중단",
        min_value=3, max_value=30, value=10,
        help="연속으로 이 횟수만큼 실패하면 봇 차단으로 판단하고 수집을 중단합니다"
    )

    run_btn = st.button("수집 시작", type="primary", use_container_width=True)

# ============================================================
# 세션 상태
# ============================================================
for key, default in [
    ('collected', False), ('df', None), ('errors', []),
    ('csv_data', None), ('csv_name', ''), ('xlsx_data', None),
    ('xlsx_name', ''), ('zip_data', None), ('zip_count', 0),
    ('zip_name', ''), ('zip_format', ''),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ============================================================
# 메인 실행
# ============================================================
if run_btn and playlist_url:

    st.session_state.collected = False

    ytdlp_ver, ytdlp_ok = check_ytdlp()
    has_ffmpeg = check_ffmpeg()

    if not ytdlp_ok:
        st.error("yt-dlp를 찾을 수 없습니다. `pip install yt-dlp`를 실행하세요.")
        st.stop()

    if not has_ffmpeg:
        st.warning("ffmpeg 미설치 — packages.txt에 `ffmpeg`를 추가하면 자막 변환 품질이 향상됩니다.")

    for d in [SUBTITLE_DIR, CONVERTED_DIR]:
        if os.path.exists(d):
            shutil.rmtree(d)
        os.makedirs(d, exist_ok=True)

    with st.status("수집 중...", expanded=True) as status:

        # ── 1단계: 플레이리스트 파싱 ──
        st.write("플레이리스트 분석 중...")
        flat_stdout, flat_stderr, flat_code = run_cmd(
            ["yt-dlp", "--flat-playlist", "--dump-json",
             "--no-warnings", "--ignore-errors", playlist_url],
            timeout=600,
        )

        if flat_code != 0 and not flat_stdout.strip():
            st.error(f"플레이리스트 파싱 실패:\n```\n{flat_stderr[:500]}\n```")
            st.stop()

        flat_entries = []
        for line in flat_stdout.strip().split('\n'):
            if line.strip():
                try:
                    flat_entries.append(json.loads(line))
                except json.JSONDecodeError:
                    continue

        video_ids = [e.get('id') or e.get('url', '') for e in flat_entries]
        video_ids = [v for v in video_ids if v]
        st.write(f"**{len(video_ids)}개** 영상 감지")

        if not video_ids:
            st.error("영상을 찾을 수 없습니다. URL을 확인하세요.")
            st.stop()

        # ── 1.5단계: 연결 테스트 ──
        st.write("YouTube 연결 테스트 중...")
        test_vid = video_ids[0]
        test_url = f"https://www.youtube.com/watch?v={test_vid}"
        test_stdout, test_stderr, test_code = run_cmd(
            ["yt-dlp", "--skip-download", "--dump-json",
             "--no-warnings", "--ignore-errors", test_url],
            timeout=30,
        )

        if test_stdout.strip():
            st.write("연결 확인 ✓")
        else:
            bot_keywords = ["Sign in", "bot", "confirm", "not a bot"]
            is_bot_block = any(kw.lower() in test_stderr.lower() for kw in bot_keywords)

            if is_bot_block:
                st.error(
                    "🚫 **YouTube가 이 네트워크의 접근을 차단하고 있습니다.**\n\n"
                    "**해결 방법:**\n"
                    "- VPN을 사용 중이라면 끄고 다시 시도\n"
                    "- 공유기를 껐다 켜서 IP 갱신 후 재시도\n"
                    "- 1~2시간 후 재시도\n"
                )
                with st.expander("yt-dlp 에러 상세"):
                    st.code(test_stderr[:1000])
                status.update(label="YouTube 봇 차단으로 중단", state="error")
                st.stop()
            else:
                st.error(f"첫 번째 영상 메타데이터 수집 실패 (code={test_code})")
                with st.expander("yt-dlp 에러 상세"):
                    st.code(test_stderr[:1000] if test_stderr else "(출력 없음)")
                st.warning("첫 번째 영상 테스트 실패. 나머지 영상으로 계속 진행합니다.")

        # ── 2단계: 개별 영상 수집 ──
        st.write(f"개별 영상 메타데이터 + 자막 수집 중... (workers={workers}, 대기={sleep_sec}초)")
        progress = st.progress(0)
        status_text = st.empty()
        full_entries = []
        errors = []
        lock = threading.Lock()
        completed_count = 0
        total = len(video_ids)

        # ★ nonlocal 대신 딕셔너리 사용
        shared = {"consecutive_fail_count": 0, "abort_flag": False}

        def process_video(idx, vid):
            """단일 영상 메타데이터 + 자막 수집"""
            if shared["abort_flag"]:
                return None, {
                    'position': idx, 'video_id': vid,
                    'error': 'Aborted (봇 차단 감지로 중단)',
                    'detail': ''
                }

            url = f"https://www.youtube.com/watch?v={vid}"
            entry = None
            error_info = None

            if sleep_sec > 0:
                time.sleep(sleep_sec)

            meta_stdout, meta_stderr, meta_code = run_cmd(
                ["yt-dlp", "--skip-download", "--dump-json",
                 "--no-warnings", "--ignore-errors", url],
                timeout=60,
            )

            if meta_stdout.strip():
                try:
                    entry = json.loads(meta_stdout.strip().split('\n')[0])
                    entry['_playlist_position'] = idx

                    with lock:
                        shared["consecutive_fail_count"] = 0

                except json.JSONDecodeError:
                    error_info = {
                        'position': idx, 'video_id': vid,
                        'error': 'JSON parse error',
                        'detail': meta_stdout[:200]
                    }
                    return entry, error_info
            else:
                bot_keywords = ["Sign in", "bot", "confirm", "not a bot"]
                is_bot = any(kw.lower() in meta_stderr.lower() for kw in bot_keywords)

                error_info = {
                    'position': idx, 'video_id': vid,
                    'error': f"No metadata (code={meta_code})",
                    'detail': meta_stderr[:300]
                }

                if is_bot:
                    with lock:
                        shared["consecutive_fail_count"] += 1
                        if shared["consecutive_fail_count"] >= max_consecutive_fails:
                            shared["abort_flag"] = True

                return entry, error_info

            # ── 자막 ──
            sub_args = [
                "yt-dlp", "--skip-download",
                "--no-warnings", "--ignore-errors",
                "-o", os.path.join(SUBTITLE_DIR, "%(id)s.%(ext)s"),
            ]

            if sub_choice == "manual_only":
                sub_args += ["--write-subs", "--no-write-auto-subs"]
            elif sub_choice == "auto_only":
                sub_args += ["--write-auto-subs"]
            elif sub_choice == "manual_first":
                sub_args += ["--write-subs", "--write-auto-subs"]
            elif sub_choice == "both":
                sub_args += ["--write-subs", "--write-auto-subs"]

            if has_ffmpeg:
                sub_args += ["--convert-subs", "srt"]

            if sub_lang.lower().strip() == "all":
                sub_args += ["--sub-langs", "all,-live_chat"]
            else:
                langs = [l.strip() for l in sub_lang.split(',') if l.strip()]
                expanded = []
                for l in langs:
                    expanded.append(l)
                    expanded.append(f"{l}-*")
                sub_args += ["--sub-langs", ','.join(expanded) + ",-live_chat"]

            sub_args.append(url)
            run_cmd(sub_args, timeout=120)

            return entry, error_info

        # ── 병렬 실행 ──
        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = {
                executor.submit(process_video, idx, vid): (idx, vid)
                for idx, vid in enumerate(video_ids, 1)
            }
            for future in as_completed(futures):
                try:
                    entry, error_info = future.result()
                except Exception as exc:
                    entry = None
                    error_info = {'position': '?', 'video_id': '?',
                                  'error': f'Future exception: {exc}'}

                with lock:
                    if entry:
                        full_entries.append(entry)
                    if error_info:
                        errors.append(error_info)
                    completed_count += 1
                    progress.progress(
                        completed_count / total,
                        text=f"[{completed_count}/{total}] 성공 {len(full_entries)} / 실패 {len(errors)}"
                    )
                    status_text.text(
                        f"처리: {completed_count}/{total} | "
                        f"성공: {len(full_entries)} | "
                        f"실패: {len(errors)} | "
                        f'연속실패: {shared["consecutive_fail_count"]}'
                    )

                if shared["abort_flag"]:
                    for f in futures:
                        f.cancel()
                    break

        if shared["abort_flag"]:
            st.warning(
                f"⚠️ **연속 {max_consecutive_fails}회 봇 차단 감지 → 수집 조기 중단**\n\n"
                f"성공한 **{len(full_entries)}개** 영상은 정상 처리됩니다.\n"
                f"나머지 영상은 1~2시간 후 재시도하거나, 공유기를 재부팅하여 IP를 갱신하세요.\n\n"
                f"**팁:** 사이드바에서 '대기 시간'을 5~10초로 늘리면 차단 확률이 줄어듭니다."
            )

        full_entries.sort(key=lambda x: x.get('_playlist_position', 0))
        progress.progress(1.0, text="수집 완료!")

        # ── 자막 수집 결과 확인 ──
        vids_with_subs, sub_file_count = count_videos_with_subs(SUBTITLE_DIR)
        st.write(f"자막 파일 **{sub_file_count}개** 수집됨 (영상 **{len(vids_with_subs)}개**)")

        if sub_file_count == 0 and full_entries:
            all_files = os.listdir(SUBTITLE_DIR) if os.path.exists(SUBTITLE_DIR) else []
            if all_files:
                st.warning(f"디렉토리에 파일 {len(all_files)}개 존재하나 크기 0:\n"
                           f"`{all_files[:5]}`")
            test_vid = full_entries[0].get('id', '')
            test_url = f"https://www.youtube.com/watch?v={test_vid}"
            test_args = [
                "yt-dlp", "--skip-download", "--write-auto-subs",
                "--sub-langs", "ko,ko-*,-live_chat",
                "--list-subs", test_url
            ]
            test_out, test_err, _ = run_cmd(test_args, timeout=30)
            with st.expander("자막 디버그 (첫 번째 영상)"):
                st.code(test_out[:1000] if test_out else "(stdout 없음)")
                st.code(test_err[:1000] if test_err else "(stderr 없음)")

        # ── 3단계: 포맷 변환 ──
        final_sub_dir = SUBTITLE_DIR
        final_sub_ext = "srt" if has_ffmpeg else "vtt"
        converted_count = 0

        actual_sub_files = []
        for fname in os.listdir(SUBTITLE_DIR):
            fpath = os.path.join(SUBTITLE_DIR, fname)
            if os.path.isfile(fpath) and os.path.getsize(fpath) > 0:
                actual_sub_files.append(fpath)
        if actual_sub_files:
            final_sub_ext = os.path.splitext(actual_sub_files[0])[1].lstrip('.')

        if output_format in ("txt", "docx") and actual_sub_files:
            st.write(f"{output_format.upper()} 변환 중...")
            if output_format == "docx":
                from docx import Document
                from docx.shared import Pt

            vid_to_files = {}
            for fpath in actual_sub_files:
                vid_from_file = os.path.basename(fpath).split('.')[0]
                vid_to_files.setdefault(vid_from_file, []).append(fpath)

            for vid_key, fpaths in vid_to_files.items():
                all_plain = []
                for fpath in fpaths:
                    lang = parse_subtitle_lang(fpath)
                    with open(fpath, 'r', encoding='utf-8', errors='replace') as f:
                        raw = f.read()
                    plain = srt_to_plain_text(raw)
                    if plain.strip():
                        if len(fpaths) > 1:
                            all_plain.append(f"[{lang}]\n{plain}")
                        else:
                            all_plain.append(plain)

                if not all_plain:
                    continue

                combined = '\n\n'.join(all_plain)
                matched_entry = next(
                    (e for e in full_entries if e.get('id') == vid_key), {}
                )
                title = matched_entry.get('title', vid_key)
                safe_name = re.sub(r'[^\w가-힣\s]', '', title)[:30].strip()

                if output_format == "txt":
                    out_path = os.path.join(CONVERTED_DIR, f"{vid_key}_{safe_name}.txt")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(f"제목: {title}\n")
                        f.write(f"영상: https://www.youtube.com/watch?v={vid_key}\n")
                        f.write(f"{'=' * 60}\n\n")
                        f.write(combined)
                    converted_count += 1

                elif output_format == "docx":
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.size = Pt(10)
                    style.paragraph_format.line_spacing = 1.5
                    doc.add_heading(title, level=1)
                    p = doc.add_paragraph()
                    p.add_run("영상: ").bold = True
                    p.add_run(f"https://www.youtube.com/watch?v={vid_key}")
                    doc.add_paragraph('─' * 40)
                    for chunk in combined.split('. '):
                        if chunk.strip():
                            doc.add_paragraph(chunk.strip() + '.')
                    out_path = os.path.join(CONVERTED_DIR, f"{vid_key}_{safe_name}.docx")
                    doc.save(out_path)
                    converted_count += 1

            final_sub_dir = CONVERTED_DIR
            final_sub_ext = output_format
            st.write(f"변환 완료: **{converted_count}개**")

        if errors:
            st.write(f"⚠️ 실패: **{len(errors)}개** 영상")

        status.update(
            label=f"수집 완료: {len(full_entries)}개 영상", state="complete"
        )

    # ── 4단계: DataFrame ──
    rows = []
    for entry in full_entries:
        vid = entry.get('id', '')
        sub_raw = read_subtitles_for_video(vid, SUBTITLE_DIR)
        sub_plain = {}
        for lang, content in sub_raw.items():
            text = srt_to_plain_text(content)
            if text.strip():
                sub_plain[lang] = text

        manual_subs = list(entry.get('subtitles', {}).keys()) if entry.get('subtitles') else []
        auto_subs = list(entry.get('automatic_captions', {}).keys()) if entry.get('automatic_captions') else []
        chapters = entry.get('chapters', [])
        chapters_str = ' | '.join(
            [f"{format_duration(ch.get('start_time', 0))} {ch.get('title', '')}"
             for ch in chapters]
        ) if chapters else ''
        thumbnails = entry.get('thumbnails', [])
        best_thumb = thumbnails[-1].get('url', '') if thumbnails else ''

        row = {
            '#': entry.get('_playlist_position', ''),
            'video_url': f"https://www.youtube.com/watch?v={vid}",
            'video_id': vid,
            'title': entry.get('title', ''),
            'description': entry.get('description', ''),
            'channel': entry.get('channel', ''),
            'channel_id': entry.get('channel_id', ''),
            'channel_url': entry.get('channel_url', ''),
            'uploader': entry.get('uploader', ''),
            'channel_follower_count': entry.get('channel_follower_count', ''),
            'upload_date': entry.get('upload_date', ''),
            'view_count': entry.get('view_count', ''),
            'like_count': entry.get('like_count', ''),
            'comment_count': entry.get('comment_count', ''),
            'duration_seconds': entry.get('duration', ''),
            'duration_readable': format_duration(entry.get('duration')),
            'categories': ', '.join(entry.get('categories') or []),
            'tags': ', '.join(entry.get('tags') or []),
            'language': entry.get('language', ''),
            'age_limit': entry.get('age_limit', ''),
            'live_status': entry.get('live_status', ''),
            'availability': entry.get('availability', ''),
            'thumbnail_url': best_thumb,
            'chapters': chapters_str,
            'manual_subtitle_langs': ', '.join(manual_subs[:30]) if manual_subs else '',
            'auto_subtitle_langs': ', '.join(auto_subs[:15]) if auto_subs else '',
            'subtitle_collected_langs': ', '.join(sub_plain.keys()),
        }
        for lang, text in sub_plain.items():
            row[f'subtitle_text_{lang}'] = text
        rows.append(row)

    df = pd.DataFrame(rows) if rows else pd.DataFrame()

    # ── 다운로드 데이터 ──
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    if not df.empty:
        st.session_state.csv_data = df.to_csv(
            index=False, encoding='utf-8-sig'
        ).encode('utf-8-sig')
        st.session_state.csv_name = f"playlist_{timestamp}.csv"

        xlsx_buf = BytesIO()
        with pd.ExcelWriter(xlsx_buf, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Videos')
            sub_cols = [c for c in df.columns if c.startswith('subtitle_text_')]
            if sub_cols:
                df[['#', 'video_id', 'title'] + sub_cols].to_excel(
                    writer, index=False, sheet_name='Subtitles'
                )
        st.session_state.xlsx_data = xlsx_buf.getvalue()
        st.session_state.xlsx_name = f"playlist_{timestamp}.xlsx"

        zip_data, zip_count = zip_directory_all(final_sub_dir, final_sub_ext)
        st.session_state.zip_data = zip_data if zip_count > 0 else None
        st.session_state.zip_count = zip_count
        st.session_state.zip_name = f"subtitles_{output_format}_{timestamp}.zip"
        st.session_state.zip_format = output_format

    st.session_state.df = df
    st.session_state.errors = errors
    st.session_state.collected = True

# ============================================================
# 결과 표시
# ============================================================
if st.session_state.collected and st.session_state.df is not None:
    df = st.session_state.df
    errors = st.session_state.errors

    if df.empty:
        st.warning("수집된 영상이 없습니다.")
        if errors:
            with st.expander(f"에러 로그 ({len(errors)}건)", expanded=True):
                st.dataframe(pd.DataFrame(errors))
        st.stop()

    sub_text_cols = [c for c in df.columns if c.startswith('subtitle_text_')]
    if sub_text_cols:
        sub_count = df[sub_text_cols].apply(
            lambda row: any(str(v).strip() not in ('', 'nan', 'None') for v in row),
            axis=1
        ).sum()
    else:
        sub_count = 0

    c1, c2, c3 = st.columns(3)
    c1.metric("총 영상", f"{len(df)}개")
    c2.metric("자막 수집", f"{sub_count}개")
    c3.metric("실패", f"{len(errors)}개")

    display_cols = ['#', 'title', 'channel', 'duration_readable',
                    'view_count', 'like_count', 'subtitle_collected_langs']
    display_cols = [c for c in display_cols if c in df.columns]
    st.dataframe(df[display_cols], use_container_width=True, height=400)

    st.subheader("다운로드")
    d1, d2, d3 = st.columns(3)

    with d1:
        if st.session_state.csv_data:
            st.markdown(make_download_link(
                st.session_state.csv_data, st.session_state.csv_name, "CSV"
            ), unsafe_allow_html=True)

    with d2:
        if st.session_state.xlsx_data:
            st.markdown(make_download_link(
                st.session_state.xlsx_data, st.session_state.xlsx_name, "XLSX"
            ), unsafe_allow_html=True)

    with d3:
        if st.session_state.zip_data:
            st.markdown(make_download_link(
                st.session_state.zip_data, st.session_state.zip_name,
                f"자막 ZIP ({st.session_state.zip_format}, {st.session_state.zip_count}개)"
            ), unsafe_allow_html=True)

    if errors:
        with st.expander(f"실패 로그 ({len(errors)}건)"):
            st.dataframe(pd.DataFrame(errors))

elif run_btn and not playlist_url:
    st.warning("플레이리스트 URL을 입력하세요.")
